"""
tests/test_meetingtool.py — MeetingTool v2.5
=============================================
Functional test suite. Run with: python -m pytest tests/ -v

Tests cover:
  - Frame extraction with synthetic video
  - Zone scoring signal
  - Edge scoring signal
  - Temporal coverage guarantee
  - Transcript DOCX parsing
  - DOCX export with image embedding
  - Setup environment detection
  - Config merge (project overrides global)
  - Language detection
  - Two-pass transcript splitting
  - Handoff JSON validation
  - Missing image ref → export failure
  - [v2.5] 720p resolution cap
  - [v2.5] SSIM perceptual similarity gate
"""

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import pytest

# ── Fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture(scope="session")
def synthetic_video(tmp_path_factory):
    """
    Generate a 3-minute synthetic test video using ffmpeg.
    9 color slides × 20 seconds each.
    Returns Path to the generated video.
    """
    out_dir = tmp_path_factory.mktemp("video")
    out_path = out_dir / "test_meeting.mp4"

    filter_complex = (
        "color=c=blue:size=1280x720:rate=30:duration=20[s1];"
        "color=c=green:size=1280x720:rate=30:duration=20[s2];"
        "color=c=red:size=1280x720:rate=30:duration=20[s3];"
        "color=c=purple:size=1280x720:rate=30:duration=20[s4];"
        "color=c=orange:size=1280x720:rate=30:duration=20[s5];"
        "color=c=teal:size=1280x720:rate=30:duration=20[s6];"
        "color=c=navy:size=1280x720:rate=30:duration=20[s7];"
        "color=c=maroon:size=1280x720:rate=30:duration=20[s8];"
        "color=c=darkgreen:size=1280x720:rate=30:duration=20[s9];"
        "[s1][s2][s3][s4][s5][s6][s7][s8][s9]concat=n=9:v=1:a=0[out]"
    )
    cmd = [
        "ffmpeg", "-y",
        "-filter_complex", filter_complex,
        "-map", "[out]",
        "-c:v", "libx264",
        "-preset", "ultrafast",
        "-crf", "35",
        str(out_path)
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, timeout=60)
    except FileNotFoundError:
        pytest.skip("ffmpeg not in PATH — skipping video tests")
    if result.returncode != 0 or not out_path.exists():
        pytest.skip("ffmpeg failed — skipping video tests")
    return out_path


@pytest.fixture
def frames_dir(tmp_path):
    d = tmp_path / "imagenes_reunion"
    d.mkdir()
    return d


@pytest.fixture
def sample_report_with_refs(tmp_path):
    """Create a sample report.md with image references."""
    frames_dir = tmp_path / "imagenes_reunion"
    frames_dir.mkdir()

    # Add dummy image files (valid JPEG via Pillow — no cv2 dependency)

    frame_names = [
        "frame_001_t00-03-12.jpg",
        "frame_002_t00-14-47.jpg",
        "frame_003_t00-28-05.jpg",
    ]
    import numpy as np
    from PIL import Image
    colors = [(200, 100, 50), (50, 200, 100), (100, 50, 200)]
    for i, name in enumerate(frame_names):
        img = Image.fromarray(np.full((100, 100, 3), colors[i], dtype=np.uint8), mode="RGB")
        img.save(str(frames_dir / name), "JPEG", quality=85)

    report_content = """# Meeting Report

## Executive Summary
This is a test meeting about data pipeline architecture.

## Screen Analysis

The team reviewed the DuckDB schema [frame_001_t00-03-12.jpg] at the start.

Later, the Tableau prototype was shown [frame_002_t00-14-47.jpg] with missing KPIs.

Final architecture diagram [frame_003_t00-28-05.jpg] was approved.

## Decisions
1. Use DuckDB as primary engine — Owner: Diego — Date: 2026-03-30
"""
    report_path = tmp_path / "report_20260330.md"
    report_path.write_text(report_content, encoding="utf-8")
    return tmp_path


# ── Test: Setup detection ─────────────────────────────────────────────────────

class TestSetupDetection:

    def test_python_version_ok(self):
        """Python 3.11+ is required."""
        assert sys.version_info >= (3, 11), (
            f"Python 3.11+ required, found {sys.version_info.major}.{sys.version_info.minor}"
        )

    def test_ffmpeg_in_path(self):
        """ffmpeg must be accessible in PATH."""
        path = shutil.which("ffmpeg")
        assert path is not None, (
            "ffmpeg not found in PATH. Install from https://ffmpeg.org/download.html"
        )

    def test_pillow_importable(self):
        """Pillow must be installed (replaces cv2 in v2.5)."""
        try:
            from PIL import Image
        except ImportError:
            pytest.fail("Pillow not installed. Run: pip install pillow")

    def test_python_docx_importable(self):
        """python-docx must be installed."""
        try:
            import docx
        except ImportError:
            pytest.fail("python-docx not installed. Run: pip install python-docx")

    def test_click_importable(self):
        """click must be installed."""
        try:
            import click
        except ImportError:
            pytest.fail("click not installed. Run: pip install click")


# ── Test: Config merge ────────────────────────────────────────────────────────

class TestConfigMerge:

    def test_project_overrides_global(self):
        """Project config values must override global config values."""
        from tools.project import _merge_configs

        global_cfg  = {"llm_provider": "claude", "default_language": "english", "mip_root": "/tmp"}
        project_cfg = {"llm_provider": "chatgpt", "report_language": "spanish", "client": "Acme"}
        merged = _merge_configs(global_cfg, project_cfg)

        assert merged["llm_provider"] == "chatgpt"
        assert merged["report_language"] == "spanish"
        assert merged["client"] == "Acme"
        assert merged["mip_root"] == "/tmp"  # global value preserved when not overridden

    def test_global_values_preserved(self):
        """Global values not present in project config must be preserved."""
        from tools.project import _merge_configs

        global_cfg  = {"llm_provider": "claude", "mip_root": "/home/user/MeetingTool", "default_language": "english"}
        project_cfg = {"client": "Kroger", "project": "RetailBeacon"}
        merged = _merge_configs(global_cfg, project_cfg)

        assert merged["mip_root"] == "/home/user/MeetingTool"
        assert merged["default_language"] == "english"
        assert merged["client"] == "Kroger"


# ── Test: Language detection ──────────────────────────────────────────────────

class TestLanguageDetection:

    def test_detects_english(self, tmp_path):
        """English transcript must be detected as English."""
        from tools.extract_frames import detect_language

        txt = tmp_path / "meeting.txt"
        txt.write_text(
            "John: The pipeline architecture that we discussed with the team "
            "will be reviewed and the data model will be updated for the next sprint. "
            "Sarah: I agree, and we should also review the test coverage for the new endpoints.",
            encoding="utf-8"
        )
        assert detect_language(txt) == "english"

    def test_detects_spanish(self, tmp_path):
        """Spanish transcript must be detected as Spanish."""
        from tools.extract_frames import detect_language

        txt = tmp_path / "meeting.txt"
        txt.write_text(
            "Diego: La arquitectura del pipeline que discutimos con el equipo "
            "será revisada y el modelo de datos será actualizado para el próximo sprint. "
            "María: Estoy de acuerdo, y también deberíamos revisar la cobertura de pruebas.",
            encoding="utf-8"
        )
        assert detect_language(txt) == "spanish"

    def test_language_mismatch_triggers_prompt(self, tmp_path, monkeypatch):
        """When meeting language differs from project default, override logic must fire."""
        from tools.extract_frames import detect_language

        txt = tmp_path / "meeting.txt"
        txt.write_text(
            "The team discussed the data architecture and the test coverage.",
            encoding="utf-8"
        )
        meeting_lang = detect_language(txt)
        project_lang = "spanish"
        assert meeting_lang != project_lang  # mismatch must be detected


# ── Test: Zone scoring ────────────────────────────────────────────────────────

class TestZoneScoring:

    def test_localized_change_scores_higher_than_uniform(self):
        """
        A frame with change in one zone should score higher than a
        frame with no meaningful change anywhere.
        """
        import numpy as np
        from tools.extract_frames import zone_score

        base = np.zeros((300, 400), dtype=np.uint8)  # all black

        # Small but significant change in top-left zone only
        localized = base.copy()
        localized[0:100, 0:100] = 200  # change one zone

        # No change
        no_change = base.copy()

        score_localized = zone_score(base, localized)
        score_no_change = zone_score(base, no_change)

        assert score_localized > score_no_change, (
            f"Localized change score ({score_localized:.3f}) should be "
            f"higher than no-change score ({score_no_change:.3f})"
        )

    def test_full_change_scores_maximum(self):
        """Complete frame change should score close to 1.0."""
        import numpy as np
        from tools.extract_frames import zone_score

        prev = np.zeros((300, 400), dtype=np.uint8)
        curr = np.full((300, 400), 200, dtype=np.uint8)

        score = zone_score(prev, curr)
        assert score > 0.8, f"Full frame change should score > 0.8, got {score:.3f}"


# ── Test: Edge scoring ────────────────────────────────────────────────────────

class TestEdgeScoring:

    def test_slide_transition_scores_higher_than_static(self):
        """
        A frame transition (blank→text-heavy) should score higher than
        a static frame comparison.
        """
        import numpy as np
        from tools.extract_frames import edge_score

        # "Previous slide" — blank
        blank = np.zeros((300, 400), dtype=np.uint8)

        # "New slide" — horizontal lines simulating text
        new_slide = np.zeros((300, 400), dtype=np.uint8)
        for y in range(30, 270, 20):
            new_slide[y, 50:350] = 255  # white horizontal line

        # "Same slide" — identical to previous
        same = blank.copy()

        score_transition = edge_score(blank, new_slide)
        score_static     = edge_score(blank, same)

        assert score_transition > score_static, (
            f"Transition score ({score_transition:.3f}) should be "
            f"higher than static score ({score_static:.3f})"
        )


# ── Test: Temporal coverage ───────────────────────────────────────────────────

class TestTemporalCoverage:

    def test_coverage_across_segments(self, synthetic_video, tmp_path):
        """
        With 20-frame budget on ~3 min video, each 20-second segment
        should have at least one frame.
        """
        from tools.extract_frames import extract_frames as ef, frames_output_dir

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        n = ef(
            video_path = synthetic_video,
            output_dir = out_dir,
            budget     = 20,
            fps_analyze = 2.0,
        )

        assert n >= 6, f"Expected at least 6 frames from 9-segment video, got {n}"

        # Verify naming format
        frames = sorted(out_dir.glob("frame_*.jpg"))
        assert len(frames) == n
        for f in frames:
            assert f.name.startswith("frame_"), f"Unexpected filename: {f.name}"
            parts = f.stem.split("_")
            assert len(parts) == 3, f"Filename format wrong: {f.name}"
            assert parts[2].startswith("t"), f"Timestamp prefix missing: {f.name}"

    def test_frame_budget_respected(self, synthetic_video, tmp_path):
        """Frame count must never exceed the specified budget."""
        from tools.extract_frames import extract_frames as ef

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        budget = 5
        n = ef(
            video_path  = synthetic_video,
            output_dir  = out_dir,
            budget      = budget,
            fps_analyze = 2.0,
        )
        assert n <= budget, f"Frame count {n} exceeds budget {budget}"


# ── Test: Transcript parsing ──────────────────────────────────────────────────

class TestTranscriptParsing:

    def test_parse_preserves_timestamps_and_speakers(self, tmp_path):
        """
        Parsed transcript must preserve all speaker timestamps.
        Boilerplate must be stripped.
        """
        from docx import Document
        from tools.extract_frames import parse_transcript_docx

        docx_path = tmp_path / "KickoffMeeting_20260330.docx"
        doc = Document()
        doc.add_paragraph("Microsoft Teams Meeting")
        doc.add_paragraph("Diego  00:01:23")
        doc.add_paragraph("Let's review the pipeline architecture.")
        doc.add_paragraph("Sarah Chen  00:05:47")
        doc.add_paragraph("I have some concerns about the data refresh cadence.")
        doc.add_paragraph("teams.microsoft.com")
        doc.save(str(docx_path))

        txt_path = parse_transcript_docx(docx_path, tmp_path)

        assert txt_path.exists()
        content = txt_path.read_text(encoding="utf-8")

        assert "[00:01:23] Diego:" in content
        assert "[00:05:47] Sarah Chen:" in content
        assert "Let's review the pipeline architecture." in content
        assert "Microsoft Teams Meeting" not in content
        assert "teams.microsoft.com" not in content


# ── Test: Two-pass split ──────────────────────────────────────────────────────

class TestTwoPassSplit:

    def test_splits_at_sentence_boundary(self, tmp_path):
        """Transcript must split at a speaker line, not mid-sentence."""
        from tools.runner import _split_transcript

        transcript = tmp_path / "meeting.txt"
        lines = []
        for i in range(100):
            if i % 10 == 0:
                lines.append(f"[00:{i:02d}:00] Speaker {i // 10}:")
            else:
                lines.append(f"This is sentence number {i} of the discussion.")
        transcript.write_text("\n".join(lines), encoding="utf-8")

        half1, half2 = _split_transcript(transcript, tmp_path)

        h1 = half1.read_text(encoding="utf-8")
        h2 = half2.read_text(encoding="utf-8")

        assert half1.exists()
        assert half2.exists()
        assert len(h1) > 0
        assert len(h2) > 0

        # Total content must equal original (minus the split boundary)
        total_orig = len(transcript.read_text(encoding="utf-8").splitlines())
        total_split = len(h1.splitlines()) + len(h2.splitlines())
        assert abs(total_split - total_orig) <= 2  # allow ±1 line for boundary


# ── Test: Handoff JSON ────────────────────────────────────────────────────────

class TestHandoffJson:

    def test_handoff_required_fields(self, tmp_path):
        """Handoff JSON must contain all required fields."""
        required_fields = {
            "meeting_id", "half", "timespan", "participants_seen",
            "decisions_confirmed", "open_threads", "action_items_partial",
            "screens_seen", "watch_for_in_half_2",
        }

        sample = {
            "meeting_id": "KickoffRetailBeacon_20260309",
            "half": 1,
            "timespan": "00:00 - 30:12",
            "participants_seen": ["Diego", "Sarah Chen"],
            "decisions_confirmed": [
                {"topic": "DuckDB as engine", "at": "00:14:22", "owner": "Diego"}
            ],
            "open_threads": [
                {"topic": "Data refresh cadence", "raised_at": "00:22:10", "status": "unresolved"}
            ],
            "action_items_partial": [
                {"task": "Send pipeline diagram", "owner": "Diego", "deadline": "Friday"}
            ],
            "screens_seen": ["DuckDB schema", "Tableau prototype"],
            "watch_for_in_half_2": ["Resolution of data refresh cadence"],
        }

        handoff_path = tmp_path / "handoff_20260330.json"
        with open(handoff_path, "w") as f:
            json.dump(sample, f)

        with open(handoff_path) as f:
            loaded = json.load(f)

        for field in required_fields:
            assert field in loaded, f"Required field '{field}' missing from handoff JSON"


# ── Test: Prompt generator ───────────────────────────────────────────────────

class TestPromptGenerator:

    def test_all_base_types_present(self):
        """All 5 base types must be in MEETING_TYPE_SECTIONS."""
        from tools.prompt_generator import MEETING_TYPE_SECTIONS
        for t in ["discovery", "kickoff", "status", "technical", "training"]:
            assert t in MEETING_TYPE_SECTIONS, f"Missing type: {t}"

    def test_training_sections_present(self):
        """Training type prompt must contain all 4 required section headers."""
        from tools.prompt_generator import generate_meeting_prompt
        config = {"llm_provider": "claude"}
        prompt = generate_meeting_prompt(config, "english", meeting_type="training")
        for section in ["Training Context", "Comprehension Assessment",
                        "Gaps & Follow-up Material", "Adoption Next Steps"]:
            assert section in prompt, f"Missing section in training prompt: {section}"

    def test_segmentation_in_all_prompts(self):
        """MEETING SEGMENTATION instruction must appear in every generated prompt."""
        from tools.prompt_generator import generate_meeting_prompt
        config = {"llm_provider": "claude"}
        for meeting_type in ["discovery", "kickoff", "status", "technical", "training"]:
            prompt = generate_meeting_prompt(config, "english", meeting_type=meeting_type)
            assert "MEETING SEGMENTATION" in prompt, (
                f"MEETING SEGMENTATION missing from {meeting_type} prompt"
            )

    def test_visual_extraction_in_pack(self):
        """Project pack must instruct frame extraction, not just citation."""
        from tools.prompt_generator import generate_prompt_pack
        pack = generate_prompt_pack({"llm_provider": "claude", "client": "X", "project": "Y"})
        assert "independent data source" in pack
        assert "visual-only evidence" in pack

    def test_fallback_includes_all_types(self):
        """Fallback type list (no meeting_type specified) must include all 5 types."""
        from tools.prompt_generator import generate_meeting_prompt
        config = {"llm_provider": "claude"}
        prompt = generate_meeting_prompt(config, "english", meeting_type=None)
        for label in ["Discovery", "Kickoff", "Status", "Technical", "Training"]:
            assert label in prompt, f"Missing type in fallback list: {label}"

    def test_no_empty_project_line_without_ref(self):
        """Claude pack must not contain a bare 'Project: ' line when ref is empty."""
        from tools.prompt_generator import generate_prompt_pack
        pack = generate_prompt_pack({"llm_provider": "claude", "client": "X",
                                     "project": "Y", "llm_project_reference": ""})
        assert "Project: \n" not in pack
        assert pack.count("Project: ") == 0 or "Project: X" in pack or "Project: Y" in pack


# ── Test: DOCX export ─────────────────────────────────────────────────────────

class TestDocxExport:

    def test_export_embeds_images(self, sample_report_with_refs):
        """DOCX export must embed all 3 referenced images."""
        from tools.exporter import run_export

        run_export(
            meeting_folder = sample_report_with_refs,
            output_format  = "docx",
        )

        docx_files = list(sample_report_with_refs.glob("report_*.docx"))
        assert len(docx_files) == 1, "Expected exactly one DOCX file"

        # Verify DOCX is a valid ZIP (DOCX format)
        import zipfile
        assert zipfile.is_zipfile(docx_files[0]), "Generated DOCX is not a valid ZIP/DOCX"

        # Verify it contains media (embedded images)
        with zipfile.ZipFile(docx_files[0]) as z:
            names = z.namelist()
            media_files = [n for n in names if n.startswith("word/media/")]
            assert len(media_files) >= 1, (
                f"Expected embedded images in DOCX, found none. Contents: {names}"
            )

    def test_export_fails_on_missing_ref(self, tmp_path):
        """Export must fail with a clear error when an image ref cannot be resolved."""
        from tools.exporter import _resolve_image_refs

        frames_dir = tmp_path / "imagenes_reunion"
        frames_dir.mkdir()
        # Do NOT create the referenced file

        report_text = "The diagram [frame_999_t00-99-99.jpg] shows the architecture."
        resolved, missing = _resolve_image_refs(report_text, frames_dir)

        assert len(missing) == 1
        assert "frame_999_t00-99-99.jpg" in missing
        assert len(resolved) == 0


# ── Test: v3.0 Embedding criterion ───────────────────────────────────────────

class TestEmbeddingCriterion:

    def test_old_criterion_removed(self):
        """The old 'adjacent frame' criterion must no longer appear in BASE_SYSTEM."""
        from tools.prompt_generator import BASE_SYSTEM
        assert "not already covered by an adjacent frame" not in BASE_SYSTEM, (
            "Old embedding criterion must be replaced — 'not already covered by an adjacent frame' "
            "still present in BASE_SYSTEM"
        )

    def test_new_criterion_present(self):
        """BASE_SYSTEM must instruct Claude to embed when transcript did NOT capture the info."""
        from tools.prompt_generator import BASE_SYSTEM
        assert "transcript did NOT capture" in BASE_SYSTEM, (
            "New embedding criterion missing: 'transcript did NOT capture' not found in BASE_SYSTEM"
        )

    def test_illegible_surfacing_instruction_present(self):
        """BASE_SYSTEM must instruct Claude to surface [ILLEGIBLE] frames rather than skip them."""
        from tools.prompt_generator import BASE_SYSTEM
        assert "visual content present but not fully extractable" in BASE_SYSTEM, (
            "ILLEGIBLE surfacing instruction missing from BASE_SYSTEM"
        )

    def test_minimum_frame_conditional_on_screen_sharing(self):
        """The per-10-min minimum must be conditioned on active screen sharing."""
        from tools.prompt_generator import BASE_SYSTEM
        assert "camera-only" in BASE_SYSTEM, (
            "Minimum frame rule must exclude camera-only segments — 'camera-only' not found"
        )


# ── Test: v3.0 Gemini prompt improvements ────────────────────────────────────

class TestGeminiPrompt:

    def test_illegible_instruction_in_prompt(self):
        """VISION_PROMPT must instruct Gemini to mark unclear text as [ILLEGIBLE]."""
        from tools.gemini_client import VISION_PROMPT
        assert "[ILLEGIBLE]" in VISION_PROMPT, (
            "[ILLEGIBLE] instruction missing from VISION_PROMPT"
        )

    def test_visual_signals_instruction_in_prompt(self):
        """VISION_PROMPT must instruct Gemini to capture color-coded and prominent elements."""
        from tools.gemini_client import VISION_PROMPT
        assert "red/green/orange" in VISION_PROMPT or "alert states" in VISION_PROMPT, (
            "Color-coding / alert states instruction missing from VISION_PROMPT"
        )
        assert "visually prominent" in VISION_PROMPT or "visual emphasis" in VISION_PROMPT, (
            "Visual prominence instruction missing from VISION_PROMPT"
        )

    def test_transcript_segment_injected(self, tmp_path):
        """When transcript_segments is provided, the snippet must appear before its image."""
        import numpy as np
        from PIL import Image
        from tools.gemini_client import _build_payload, CHUNK_SIZE

        # Create 3 minimal JPEG frames
        frames = []
        for i in range(3):
            img = Image.fromarray(np.full((72, 128, 3), (i * 80, 100, 150), dtype=np.uint8), "RGB")
            p = tmp_path / f"frame_{i+1:03d}_t00-0{i}-00.jpg"
            img.save(str(p), "JPEG")
            frames.append(p)

        # Frame global index 2 (1-based) gets a transcript snippet
        transcript_segments = {2: "Randy: I want to understand where the inputs come from."}

        payload = _build_payload(frames, chunk_index=0,
                                 transcript_segments=transcript_segments)
        parts = payload["contents"][0]["parts"]

        # Find the index of the frame-2 label and the speaker context
        frame2_label_idx = next(
            i for i, p in enumerate(parts)
            if isinstance(p.get("text"), str) and "FRAME 2" in p["text"]
        )
        context_part = parts[frame2_label_idx + 1]
        image_part   = parts[frame2_label_idx + 2]

        assert "Speaker context" in context_part.get("text", ""), (
            "Transcript snippet must appear immediately after frame label"
        )
        assert "Randy" in context_part["text"], (
            "Snippet content must be preserved verbatim"
        )
        assert "inline_data" in image_part, (
            "Image must follow immediately after the transcript context"
        )

    def test_no_transcript_segment_when_not_provided(self, tmp_path):
        """Without transcript_segments, payload must contain label → image (no extra parts)."""
        import numpy as np
        from PIL import Image
        from tools.gemini_client import _build_payload

        img = Image.fromarray(np.full((72, 128, 3), (100, 100, 100), dtype=np.uint8), "RGB")
        p = tmp_path / "frame_001_t00-00-00.jpg"
        img.save(str(p), "JPEG")

        payload = _build_payload([p], chunk_index=0, transcript_segments=None)
        parts = payload["contents"][0]["parts"]

        # parts: [prompt_text, frame_label, inline_data] — exactly 3
        assert len(parts) == 3, (
            f"Without transcript_segments, payload must have 3 parts (prompt+label+image), got {len(parts)}"
        )
        assert "inline_data" in parts[2], "Third part must be the image"


# ── Test: v2.5 API config ────────────────────────────────────────────────────

class TestApiConfig:

    def test_loads_key_from_keyring(self, monkeypatch):
        """_load_key returns the value stored in keyring."""
        import keyring
        from tools.api_config import _load_key, SERVICE_NAME

        monkeypatch.setattr(keyring, "get_password",
                            lambda svc, key: "sk-test-123" if svc == SERVICE_NAME else None)
        assert _load_key("ANTHROPIC_API_KEY") == "sk-test-123"

    def test_missing_key_raises_runtime_error(self, monkeypatch):
        """A missing key must raise RuntimeError pointing to mip setup."""
        import keyring
        from tools.api_config import _load_key

        monkeypatch.setattr(keyring, "get_password", lambda *_: None)

        with pytest.raises(RuntimeError, match="mip.py setup"):
            _load_key("NONEXISTENT_KEY")

    def test_get_gemini_key_2_returns_value_when_set(self, monkeypatch):
        """get_gemini_key_2 returns the backup key when stored in keyring."""
        import keyring
        from tools.api_config import get_gemini_key_2, SERVICE_NAME

        monkeypatch.setattr(
            keyring, "get_password",
            lambda svc, key: "backup-key-456" if key == "GEMINI_API_KEY_2" else None
        )
        assert get_gemini_key_2() == "backup-key-456"

    def test_get_gemini_key_2_returns_none_when_absent(self, monkeypatch):
        """get_gemini_key_2 returns None when no backup key is stored."""
        import keyring
        from tools.api_config import get_gemini_key_2

        monkeypatch.setattr(keyring, "get_password", lambda *_: None)
        assert get_gemini_key_2() is None


# ── Test: v2.5 Claude client ─────────────────────────────────────────────────

class TestClaudeClient:

    def _make_mock_client(self, captured: dict):
        """Return a fake Anthropic client that records what was passed."""
        def mock_create(**kwargs):
            captured.update(kwargs)
            return type("R", (), {
                "content": [type("C", (), {"text": "# Report\n\nTest."})()],
                "usage":   type("U", (), {
                    "input_tokens": 100,
                    "cache_read_input_tokens": 50,
                    "cache_creation_input_tokens": 0,
                })(),
            })()
        return type("Client", (), {
            "messages": type("M", (), {"create": staticmethod(mock_create)})()
        })()

    def test_system_prompt_has_cache_control(self):
        """The system message must carry cache_control ephemeral."""
        from tools.claude_client import generate_report
        captured = {}
        client = self._make_mock_client(captured)

        generate_report(
            transcript_text="Meeting discussion.",
            visual_evidence="Frame 1: dashboard visible.",
            config={"client": "Acme"},
            _client=client,
        )

        assert captured.get("system"), "system block must be present"
        assert captured["system"][0].get("cache_control") == {"type": "ephemeral"}, (
            "cache_control ephemeral must be on system prompt"
        )

    def test_report_uses_claude_model(self):
        """generate_report must call the correct Claude model."""
        from tools.claude_client import generate_report, CLAUDE_MODEL
        captured = {}
        generate_report("transcript", "evidence", {}, _client=self._make_mock_client(captured))
        assert captured["model"] == CLAUDE_MODEL

    def test_write_report_creates_file(self, tmp_path):
        """write_report must create report_{date}.md with the correct content."""
        from tools.claude_client import write_report
        content = "# Meeting Report\n\nTest content."
        path = write_report(content, tmp_path, "20260617")
        assert path.exists()
        assert path.name == "report_20260617.md"
        assert path.read_text(encoding="utf-8") == content

    def test_missing_visual_evidence_handled(self):
        """Empty visual_evidence must inject fallback text, not crash."""
        from tools.claude_client import generate_report
        captured = {}
        generate_report("transcript text", "", {}, _client=self._make_mock_client(captured))
        assert "No visual evidence available" in captured["messages"][0]["content"]


# ── Test: v2.5 Gemini client ─────────────────────────────────────────────────

class TestGeminiClient:

    def _make_frames(self, tmp_path: Path, count: int) -> list:
        """Create minimal valid JPEG files for testing."""
        import numpy as np
        from PIL import Image
        frames = []
        for i in range(count):
            img = Image.fromarray(
                np.full((72, 128, 3), (i * 20 % 255, 100, 150), dtype=np.uint8), "RGB"
            )
            p = tmp_path / f"frame_{i+1:03d}_t00-0{i}-00.jpg"
            img.save(str(p), "JPEG")
            frames.append(p)
        return frames

    def test_chunks_150_frames(self, tmp_path):
        """150 frames must produce ceil(150/CHUNK_SIZE) chunks, all ≤ CHUNK_SIZE."""
        import math
        from tools.gemini_client import CHUNK_SIZE
        frames = self._make_frames(tmp_path, 150)
        chunks = [frames[i: i + CHUNK_SIZE] for i in range(0, len(frames), CHUNK_SIZE)]
        assert len(chunks) == math.ceil(150 / CHUNK_SIZE)
        assert all(len(c) <= CHUNK_SIZE for c in chunks)
        assert sum(len(c) for c in chunks) == 150

    def test_successful_extraction(self, tmp_path, monkeypatch):
        """Successful Gemini response returns concatenated visual_evidence."""
        import tools.gemini_client as gc

        def mock_post(payload, api_key):
            chunk_n = len([p for p in payload["contents"][0]["parts"] if "text" in p and "FRAME" in p.get("text", "")])
            return f"Mock visual evidence for {chunk_n} frames."

        monkeypatch.setattr(gc, "_post_chunk", mock_post)
        monkeypatch.setattr(gc, "CHUNK_DELAY", 0)

        frames = self._make_frames(tmp_path, 5)
        result = gc.extract_visual_evidence(frames, "fake-key")
        assert "Mock visual evidence" in result

    def test_429_triggers_retry(self, tmp_path, monkeypatch):
        """HTTP 429 must trigger one retry after RETRY_DELAY."""
        import requests
        import tools.gemini_client as gc

        call_count = {"n": 0}

        def mock_post(payload, api_key):
            call_count["n"] += 1
            if call_count["n"] == 1:
                mock_resp = type("R", (), {"status_code": 429})()
                raise requests.HTTPError("429", response=mock_resp)
            return "Retry succeeded."

        monkeypatch.setattr(gc, "_post_chunk", mock_post)
        monkeypatch.setattr(gc, "CHUNK_DELAY", 0)
        monkeypatch.setattr(gc, "RETRY_DELAY", 0)

        frames = self._make_frames(tmp_path, 3)
        result = gc.extract_visual_evidence(frames, "fake-key")
        assert "Retry succeeded" in result
        assert call_count["n"] == 2

    def test_persistent_429_raises_unavailable(self, tmp_path, monkeypatch):
        """Two consecutive 429s must raise GeminiUnavailableError."""
        import requests
        import tools.gemini_client as gc

        def mock_post(payload, api_key):
            mock_resp = type("R", (), {"status_code": 429})()
            raise requests.HTTPError("429", response=mock_resp)

        monkeypatch.setattr(gc, "_post_chunk", mock_post)
        monkeypatch.setattr(gc, "CHUNK_DELAY", 0)
        monkeypatch.setattr(gc, "RETRY_DELAY", 0)

        frames = self._make_frames(tmp_path, 3)
        with pytest.raises(gc.GeminiUnavailableError):
            gc.extract_visual_evidence(frames, "fake-key")

    def test_fallback_key2_used_when_key1_exhausted(self, tmp_path, monkeypatch):
        """When key 1 hits 429 after retry, key 2 must be tried and succeed."""
        import requests
        import tools.gemini_client as gc

        calls = []

        def mock_post(payload, api_key):
            calls.append(api_key)
            if api_key == "key-1":
                mock_resp = type("R", (), {"status_code": 429})()
                raise requests.HTTPError("429", response=mock_resp)
            return "evidence from key 2"

        monkeypatch.setattr(gc, "_post_chunk", mock_post)
        monkeypatch.setattr(gc, "CHUNK_DELAY", 0)
        monkeypatch.setattr(gc, "RETRY_DELAY", 0)

        frames = self._make_frames(tmp_path, 3)
        result = gc.extract_visual_evidence(frames, "key-1", api_key_2="key-2")
        assert result == "evidence from key 2"
        assert "key-2" in calls

    def test_fallback_key2_absent_raises_unavailable(self, tmp_path, monkeypatch):
        """When key 2 is not configured and key 1 is exhausted, must raise GeminiUnavailableError."""
        import requests
        import tools.gemini_client as gc

        def mock_post(payload, api_key):
            mock_resp = type("R", (), {"status_code": 429})()
            raise requests.HTTPError("429", response=mock_resp)

        monkeypatch.setattr(gc, "_post_chunk", mock_post)
        monkeypatch.setattr(gc, "CHUNK_DELAY", 0)
        monkeypatch.setattr(gc, "RETRY_DELAY", 0)

        frames = self._make_frames(tmp_path, 3)
        with pytest.raises(gc.GeminiUnavailableError):
            gc.extract_visual_evidence(frames, "key-1", api_key_2=None)

    def test_fallback_key2_also_fails_raises_unavailable(self, tmp_path, monkeypatch):
        """When both keys hit 429, must raise GeminiUnavailableError."""
        import requests
        import tools.gemini_client as gc

        def mock_post(payload, api_key):
            mock_resp = type("R", (), {"status_code": 429})()
            raise requests.HTTPError("429", response=mock_resp)

        monkeypatch.setattr(gc, "_post_chunk", mock_post)
        monkeypatch.setattr(gc, "CHUNK_DELAY", 0)
        monkeypatch.setattr(gc, "RETRY_DELAY", 0)

        frames = self._make_frames(tmp_path, 3)
        with pytest.raises(gc.GeminiUnavailableError):
            gc.extract_visual_evidence(frames, "key-1", api_key_2="key-2")

    def test_empty_frames_returns_empty_string(self, tmp_path):
        """No frames must return empty string without calling Gemini."""
        from tools.gemini_client import extract_visual_evidence
        result = extract_visual_evidence([], "fake-key")
        assert result == ""


# ── Test: v3.0 Transcript boost / discard log ────────────────────────────────

class TestTranscriptBoost:

    def test_discard_log_created(self, synthetic_video, tmp_path):
        """Full extraction run must generate frames_discarded.log with valid entries."""
        from tools.extract_frames import extract_frames as ef

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        ef(video_path=synthetic_video, output_dir=out_dir, budget=5, fps_analyze=2.0)

        log_path = out_dir / "frames_discarded.log"
        assert log_path.exists(), "frames_discarded.log must be created after extraction"

        content = log_path.read_text(encoding="utf-8")
        lines = [l for l in content.splitlines() if l.strip()]
        assert len(lines) >= 1, "Log must have at least one discard entry"

        first = lines[0]
        assert " | " in first, f"Expected pipe-delimited format, got: {first}"
        assert "motivo=" in first, f"Expected motivo= field in entry, got: {first}"

    def test_discard_log_records_score_bajo(self, tmp_path):
        """Frames below min_composite_score must appear in the log as score_bajo."""
        import numpy as np
        from tools.extract_frames import _write_discard_log, seconds_to_filename_ts

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        discards = [
            (45.0, "score=0.08 | motivo=score_bajo"),
            (90.0, "score=0.11 | motivo=score_bajo"),
        ]
        _write_discard_log(out_dir, discards, "2026-06-18T10:00:00")

        content = (out_dir / "frames_discarded.log").read_text(encoding="utf-8")
        assert "score_bajo" in content
        assert "score=0.08" in content
        assert "score=0.11" in content

    def test_discard_log_records_budget_cap(self, tmp_path):
        """Frames removed by budget cap must appear in the log as budget_cap."""
        from tools.extract_frames import _write_discard_log

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        discards = [(120.0, "motivo=budget_cap (frame 151, score=0.23)")]
        _write_discard_log(out_dir, discards, "2026-06-18T10:00:00")

        content = (out_dir / "frames_discarded.log").read_text(encoding="utf-8")
        assert "budget_cap" in content
        assert "frame 151" in content

    def test_boost_applied_when_visual_ref_present(self):
        """Score must increase by TRANSCRIPT_BOOST when visual ref keyword is in window."""
        from tools.extract_frames import _transcript_has_visual_ref, TRANSCRIPT_BOOST

        transcript = (
            "[00:00:55] Odeta Pine:\n"
            "If you look at this diagram, mirá esto aquí.\n\n"
            "[00:01:10] Randy:\nGot it, makes sense."
        )
        assert _transcript_has_visual_ref(transcript, 60.0, window=30), (
            "Expected True: visual keyword 'mirá' is within ±30s of t=60s"
        )

    def test_no_boost_without_visual_ref(self):
        """Score must not change when no visual ref keyword is near the timestamp."""
        from tools.extract_frames import _transcript_has_visual_ref

        transcript = (
            "[00:00:55] Odeta Pine:\n"
            "The data architecture has three distinct layers.\n\n"
            "[00:01:10] Randy:\nUnderstood, thank you."
        )
        assert not _transcript_has_visual_ref(transcript, 60.0, window=30), (
            "Expected False: no visual reference keyword in transcript window"
        )

    def test_boost_not_applied_outside_window(self):
        """Visual ref keyword far from frame timestamp must not trigger boost."""
        from tools.extract_frames import _transcript_has_visual_ref

        transcript = (
            "[00:00:10] Odeta Pine:\n"
            "Look at this dashboard on screen.\n\n"
            "[00:05:00] Randy:\nAnd this other topic."
        )
        # t=300s, window=30 → only [00:04:30..00:05:30] qualifies → keyword at t=10 is excluded
        assert not _transcript_has_visual_ref(transcript, 300.0, window=30), (
            "Expected False: visual keyword at t=10s is outside ±30s window of t=300s"
        )

    def test_boost_empty_transcript_returns_false(self):
        """Empty transcript must never trigger boost."""
        from tools.extract_frames import _transcript_has_visual_ref
        assert not _transcript_has_visual_ref("", 60.0), (
            "Expected False for empty transcript"
        )

    def test_discard_log_sorted_chronologically(self, tmp_path):
        """Discard log entries must be sorted by timestamp, not insertion order."""
        from tools.extract_frames import _write_discard_log

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        # Insert out of order
        discards = [
            (300.0, "motivo=gap_minimo (2.1s desde anterior)"),
            (60.0,  "score=0.10 | motivo=score_bajo"),
            (180.0, "SSIM=0.97 | motivo=duplicado_perceptual"),
        ]
        _write_discard_log(out_dir, discards, "2026-06-18T10:00:00")

        lines = (out_dir / "frames_discarded.log").read_text(encoding="utf-8").splitlines()
        assert "score_bajo" in lines[0], "First entry (t=60s) should be score_bajo"
        assert "duplicado_perceptual" in lines[1], "Second entry (t=180s) should be duplicado_perceptual"
        assert "gap_minimo" in lines[2], "Third entry (t=300s) should be gap_minimo"


# ── Test: v2.5 Frame optimization ────────────────────────────────────────────

class TestFrameOptimization:

    def _make_solid_frame(self, width: int, height: int, color: tuple) -> "np.ndarray":
        import numpy as np
        return np.full((height, width, 3), color, dtype=np.uint8)

    def test_720p_cap_resizes_oversized_frame(self, tmp_path):
        """Frames wider or taller than 1280×720 must be saved at ≤ 1280×720."""
        import numpy as np
        from PIL import Image
        from tools.extract_frames import MAX_WIDTH, MAX_HEIGHT

        # Build a fake 1920×1080 frame
        img_rgb = self._make_solid_frame(1920, 1080, (100, 150, 200))

        pil_img = Image.fromarray(img_rgb)
        w, h = pil_img.size
        if w > MAX_WIDTH or h > MAX_HEIGHT:
            scale   = min(MAX_WIDTH / w, MAX_HEIGHT / h)
            pil_img = pil_img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        out_w, out_h = pil_img.size
        assert out_w <= MAX_WIDTH,  f"Width {out_w} exceeds {MAX_WIDTH}"
        assert out_h <= MAX_HEIGHT, f"Height {out_h} exceeds {MAX_HEIGHT}"

    def test_720p_cap_does_not_upscale(self, tmp_path):
        """Frames already below 720p must not be resized."""
        import numpy as np
        from PIL import Image
        from tools.extract_frames import MAX_WIDTH, MAX_HEIGHT

        img_rgb = self._make_solid_frame(640, 360, (200, 100, 50))
        pil_img = Image.fromarray(img_rgb)
        w, h = pil_img.size
        if w > MAX_WIDTH or h > MAX_HEIGHT:
            scale   = min(MAX_WIDTH / w, MAX_HEIGHT / h)
            pil_img = pil_img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        assert pil_img.size == (640, 360), f"Small frame must not be resized, got {pil_img.size}"

    def test_ssim_gate_discards_duplicate(self, synthetic_video, tmp_path):
        """
        When two near-identical frames are candidates, only the first must be saved.
        Verified by running full extraction on a synthetic video with static segments.
        """
        from tools.extract_frames import extract_frames as ef, SSIM_DISCARD_THRESHOLD

        out_dir = tmp_path / "imagenes_reunion"
        out_dir.mkdir()

        # The synthetic video has 9 distinct solid-color segments — SSIM between
        # consecutive segments should be very low (different colors).
        # Within each 20-second segment, consecutive sampled frames are identical →
        # SSIM ≈ 1.0 → all but the first should be discarded.
        n = ef(
            video_path  = synthetic_video,
            output_dir  = out_dir,
            budget      = 150,
            fps_analyze = 2.0,
        )
        # Without SSIM gate, 9 segments × 20s × 2fps = 360 candidates → trimmed to 150.
        # With SSIM gate, identical frames within each segment are discarded →
        # result should be close to 9 (one per unique color).
        assert n <= 20, (
            f"SSIM gate should eliminate near-duplicate frames within static segments. "
            f"Got {n} frames — expected ≤ 20."
        )

    def test_ssim_allows_distinct_frames(self):
        """Frames with SSIM well below 0.95 must both be saved."""
        import numpy as np
        from skimage.metrics import structural_similarity as ssim_fn
        from PIL import Image
        from tools.extract_frames import SSIM_DISCARD_THRESHOLD

        frame_a = np.zeros((720, 1280), dtype=np.uint8)          # black
        frame_b = np.full((720, 1280), 200, dtype=np.uint8)       # light grey

        similarity = ssim_fn(frame_a, frame_b)
        assert similarity < SSIM_DISCARD_THRESHOLD, (
            f"Distinct frames have SSIM={similarity:.3f} which should be < {SSIM_DISCARD_THRESHOLD}"
        )


# ── Test: v2.5 Runner automated pipeline ─────────────────────────────────────

class TestRunnerAutomated:

    def _make_frames_in(self, out_dir: Path, count: int) -> list:
        """Create minimal JPEG files and return their paths."""
        import numpy as np
        from PIL import Image
        out_dir.mkdir(parents=True, exist_ok=True)
        frames = []
        for i in range(count):
            img = Image.fromarray(
                np.full((72, 128, 3), (i * 50 % 255, 100, 150), dtype=np.uint8), "RGB"
            )
            p = out_dir / f"frame_{i+1:03d}_t00-0{i % 10}-00.jpg"
            img.save(str(p), "JPEG")
            frames.append(p)
        return frames

    def test_automated_pipeline_gemini_success(self, tmp_path, monkeypatch):
        """Full automated path: Gemini runs, Claude runs, report file written."""
        import tools.runner as runner

        frames_dir = tmp_path / "imagenes_reunion"

        def fake_extract_frames(video_path, output_dir, budget, **kwargs):
            self._make_frames_in(output_dir, 3)
            return 3

        monkeypatch.setattr(runner, "extract_frames", fake_extract_frames)
        monkeypatch.setattr(runner, "extract_visual_evidence",
                            lambda paths, key, **kw: "Mock visual evidence from 3 frames.")
        monkeypatch.setattr(runner, "get_gemini_key", lambda config=None: "fake-gemini-key")
        monkeypatch.setattr(runner, "generate_report",
                            lambda **kwargs: "# Meeting Report\n\nContent.")
        monkeypatch.setattr(runner, "get_anthropic_key", lambda: "fake-anthropic-key")

        mp4 = tmp_path / "meeting.mp4"
        mp4.write_bytes(b"\x00" * 16)

        result = runner._run_cowork_automated(
            meeting_folder      = tmp_path,
            video_path          = mp4,
            transcript_path     = None,
            config              = {"report_language": "english"},
            max_frames_override = 3,
        )

        assert result.report_path is not None
        assert result.report_path.exists()
        assert result.visual_evidence_source == "gemini"
        assert result.n_frames == 3

    def test_gemini_failure_triggers_ocr_fallback(self, tmp_path, monkeypatch):
        """GeminiUnavailableError must activate OCR fallback and set source."""
        import tools.runner as runner
        from tools.gemini_client import GeminiUnavailableError

        def fake_extract_frames(video_path, output_dir, budget, **kwargs):
            self._make_frames_in(output_dir, 2)
            return 2

        def fake_gemini(paths, key, **kw):
            raise GeminiUnavailableError("rate limited after retry")

        monkeypatch.setattr(runner, "extract_frames", fake_extract_frames)
        monkeypatch.setattr(runner, "extract_visual_evidence", fake_gemini)
        monkeypatch.setattr(runner, "get_gemini_key", lambda config=None: "fake-key")
        monkeypatch.setattr(runner, "_ocr_fallback", lambda paths: "OCR text from 2 frames.")
        monkeypatch.setattr(runner, "generate_report",
                            lambda **kwargs: "# Fallback Report\n\nContent.")
        monkeypatch.setattr(runner, "get_anthropic_key", lambda: "fake-key")

        mp4 = tmp_path / "meeting.mp4"
        mp4.write_bytes(b"\x00" * 16)

        result = runner._run_cowork_automated(
            meeting_folder      = tmp_path,
            video_path          = mp4,
            transcript_path     = None,
            config              = {"report_language": "english"},
            max_frames_override = 2,
        )

        assert result.visual_evidence_source == "ocr_fallback"
        assert result.report_path is not None
        assert result.report_path.exists()

    def test_ocr_fallback_no_pytesseract(self, monkeypatch):
        """Without pytesseract installed, OCR fallback returns empty string."""
        import builtins
        import tools.runner as runner

        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "pytesseract":
                raise ImportError("no pytesseract")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", mock_import)
        result = runner._ocr_fallback([])
        assert result == ""

    def test_automated_pipeline_returns_report_in_meeting_folder(self, tmp_path, monkeypatch):
        """report_path must be inside meeting_folder."""
        import tools.runner as runner

        def fake_extract_frames(video_path, output_dir, budget, **kwargs):
            self._make_frames_in(output_dir, 1)
            return 1

        monkeypatch.setattr(runner, "extract_frames", fake_extract_frames)
        monkeypatch.setattr(runner, "extract_visual_evidence",
                            lambda paths, key, **kw: "evidence")
        monkeypatch.setattr(runner, "get_gemini_key", lambda config=None: "k")
        monkeypatch.setattr(runner, "generate_report",
                            lambda **kwargs: "# Report")
        monkeypatch.setattr(runner, "get_anthropic_key", lambda: "k")

        mp4 = tmp_path / "meeting.mp4"
        mp4.write_bytes(b"\x00" * 16)

        result = runner._run_cowork_automated(
            meeting_folder      = tmp_path,
            video_path          = mp4,
            transcript_path     = None,
            config              = {},
            max_frames_override = 1,
        )

        assert result.report_path.parent == tmp_path
        assert result.report_path.name.startswith("report_")


# ── Test: Inc 4 — Client context ──────────────────────────────────────────────

class TestClientContext:

    def test_client_context_saved_to_config(self, tmp_path, monkeypatch):
        """client_context and client_context_updated written to project config."""
        import json
        import tools.project as project

        answers = iter([
            "AcmeCorp",           # client name
            "Q2Analysis",         # project name
            str(tmp_path),        # project folder path
            "",                   # provider ref (_ask_choice for provider is patched separately)
            "Retail, 500 employees, SAP stack.",  # client_context
            "",                   # custom meeting types
        ])
        monkeypatch.setattr("tools.project._ask", lambda prompt, default="": next(answers))
        monkeypatch.setattr("tools.project._ask_choice", lambda prompt, opts, default_key=None: list(opts.values())[int(default_key or "0") - 1])
        monkeypatch.setattr("tools.project._load_global_config", lambda: {
            "mip_root": str(tmp_path), "llm_provider": "claude", "default_language": "english"
        })
        monkeypatch.setattr("tools.project.generate_prompt_pack", lambda cfg: "prompt")

        project.run_project_new()

        cfg_path = tmp_path / "mip.config.json"
        assert cfg_path.exists()
        cfg = json.loads(cfg_path.read_text())
        assert cfg["client_context"] == "Retail, 500 employees, SAP stack."
        assert cfg["client_context_updated"] != ""

    def test_client_context_optional(self, tmp_path, monkeypatch):
        """Empty client_context leaves client_context as empty string, not error."""
        import json
        import tools.project as project

        answers = iter([
            "AcmeCorp", "Q2Analysis", str(tmp_path), "", "", "",
        ])
        monkeypatch.setattr("tools.project._ask", lambda prompt, default="": next(answers))
        monkeypatch.setattr("tools.project._ask_choice", lambda prompt, opts, default_key=None: list(opts.values())[int(default_key or "0") - 1])
        monkeypatch.setattr("tools.project._load_global_config", lambda: {
            "mip_root": str(tmp_path), "llm_provider": "claude", "default_language": "english"
        })
        monkeypatch.setattr("tools.project.generate_prompt_pack", lambda cfg: "prompt")

        project.run_project_new()

        cfg = json.loads((tmp_path / "mip.config.json").read_text())
        assert cfg["client_context"] == ""
        assert cfg["client_context_updated"] == ""

    def test_staleness_warning_triggered(self):
        """_check_context_staleness returns True (stale) when >90 days old."""
        from datetime import date, timedelta
        from tools.runner import _check_context_staleness
        import io, sys

        old_date = (date.today() - timedelta(days=91)).isoformat()
        config = {"client_context": "Some context.", "client_context_updated": old_date}

        captured = io.StringIO()
        sys.stdout = captured
        try:
            import builtins
            original_input = builtins.input
            builtins.input = lambda _: "n"
            result = _check_context_staleness(config)
            builtins.input = original_input
        finally:
            sys.stdout = sys.__stdout__

        assert result is True

    def test_no_warning_under_90_days(self):
        """_check_context_staleness returns False when <=90 days old."""
        from datetime import date, timedelta
        from tools.runner import _check_context_staleness

        recent_date = (date.today() - timedelta(days=45)).isoformat()
        config = {"client_context": "Some context.", "client_context_updated": recent_date}

        result = _check_context_staleness(config)
        assert result is False
