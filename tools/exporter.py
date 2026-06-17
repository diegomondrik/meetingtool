"""
tools/exporter.py — MeetingTool v2.0
======================================
Converts report.md to report.docx with embedded images.
Resolves [frame_XXX_tHH-MM-SS.jpg] references from imagenes_reunion folder.
Fails with a clear error listing all missing refs — never silently skips.

Run via: mip export --path <meeting_folder>
"""

import re
from pathlib import Path

from tools.installer import _ok, _warn, _err


# ── Regex for image refs ──────────────────────────────────────────────────────

# Format A: exact filename  [frame_017_t00-13-03.jpg]
_IMAGE_REF = re.compile(r'\[frame_\d+_t\d{2}-\d{2}-\d{2}\.jpg\]')

# Format B: single frame, Claude report style  [frame_017, t00:13:03]
# Resolves by frame number only (glob), ignoring the timestamp.
_IMAGE_REF_ALT = re.compile(r'\[frame_(\d+),\s*t\d{2}:\d{2}:\d{2}\]')

# Format C: frame range, Claude report style  [frames_019–024, t00:13:37–t00:17:57]
# en-dash (–) or hyphen (-) between frame numbers; expands to all frames in range.
_IMAGE_REF_RANGE = re.compile(
    r'\[frames_(\d+)[–\-](\d+),\s*t\d{2}:\d{2}:\d{2}[–\-]t\d{2}:\d{2}:\d{2}\]'
)

_EXPORT_TRIGGERS = {"send", "deliver", "client", "final"}


# ── Report discovery ──────────────────────────────────────────────────────────

def _find_report_md(meeting_folder: Path) -> Path | None:
    candidates = sorted(meeting_folder.glob("report_*.md"), reverse=True)
    if candidates:
        return candidates[0]
    return None


# ── Image ref resolution ──────────────────────────────────────────────────────

def _resolve_image_refs(
    report_text: str,
    frames_dir: Path,
) -> tuple[list[tuple[str, Path]], list[str]]:
    """
    Find all image refs in the report and resolve them to actual file paths.
    Supports three formats:
      A) [frame_017_t00-13-03.jpg]           — exact filename
      B) [frame_017, t00:13:03]              — single frame, Claude style
      C) [frames_019–024, t00:13:37–t00:17:57] — frame range, Claude style
    One (ref_string, file_path) tuple is added per resolved file, so a range
    ref produces multiple tuples (one per frame in the range).
    Returns:
        resolved:  list of (ref_string, file_path)
        missing:   list of descriptions for refs that could not be resolved
    """
    seen     = set()
    resolved = []
    missing  = []

    # Format A: exact filename refs
    for ref in _IMAGE_REF.findall(report_text):
        if ref in seen:
            continue
        seen.add(ref)
        file_path = frames_dir / ref.strip("[]")
        if file_path.exists():
            resolved.append((ref, file_path))
        else:
            missing.append(ref.strip("[]"))

    # Format B: [frame_NNN, t...] — resolve by frame number only (glob)
    for m in _IMAGE_REF_ALT.finditer(report_text):
        ref = m.group(0)
        if ref in seen:
            continue
        seen.add(ref)
        num        = m.group(1).zfill(3)
        candidates = sorted(frames_dir.glob(f"frame_{num}_t*.jpg"))
        if candidates:
            resolved.append((ref, candidates[0]))
        else:
            missing.append(f"frame_{num}_t*.jpg")

    # Format C: [frames_NNN–MMM, t...–t...] — expand range, one tuple per frame
    for m in _IMAGE_REF_RANGE.finditer(report_text):
        ref = m.group(0)
        if ref in seen:
            continue
        seen.add(ref)
        start, end = int(m.group(1)), int(m.group(2))
        found_any  = False
        for n in range(start, end + 1):
            candidates = sorted(frames_dir.glob(f"frame_{str(n).zfill(3)}_t*.jpg"))
            if candidates:
                resolved.append((ref, candidates[0]))
                found_any = True
        if not found_any:
            missing.append(f"frames_{m.group(1)}–{m.group(2)}_t*.jpg")

    return resolved, missing


# ── Markdown → DOCX conversion ────────────────────────────────────────────────

def _md_to_docx(
    report_text: str,
    resolved_refs: list[tuple[str, Path]],
    output_path: Path,
    meeting_folder: Path,
):
    """
    Convert Markdown report to DOCX with embedded images.
    Images are inserted inline at the point of each [frame_XXX_...] reference.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed — run: pip install python-docx")

    # Build a lookup: ref_string → [path, ...] (list because ranges map to multiple frames)
    ref_map: dict = {}
    for ref, path in resolved_refs:
        ref_map.setdefault(ref, []).append(path)

    doc = Document()

    # ── Document styles ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── Page margins (1 inch) ──
    from docx.shared import Inches as In
    for section in doc.sections:
        section.top_margin    = In(1)
        section.bottom_margin = In(1)
        section.left_margin   = In(1)
        section.right_margin  = In(1)

    # ── Parse and render markdown lines ──
    lines = report_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "___", "***"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            content = line[2:].strip()
            content, ref_paths = _resolve_inline_refs(content, ref_map)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_content(p, content, ref_paths, doc)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\. ", line):
            content = re.sub(r"^\d+\. ", "", line).strip()
            content, ref_paths = _resolve_inline_refs(content, ref_map)
            p = doc.add_paragraph(style="List Number")
            _add_inline_content(p, content, ref_paths, doc)
            i += 1
            continue

        # Table (detect by pipe characters)
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _render_table(doc, table_lines)
            continue

        # Empty line
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph (may contain image refs)
        content, ref_paths = _resolve_inline_refs(line, ref_map)
        p = doc.add_paragraph()
        _add_inline_content(p, content, ref_paths, doc)
        i += 1

    doc.save(str(output_path))


def _resolve_inline_refs(text: str, ref_map: dict) -> tuple[str, list]:
    """
    Find image refs in a text line and return (cleaned_text, list_of_image_paths).
    Image refs are replaced with a single placeholder; ref_map[ref] is a list of
    paths so one placeholder can expand to multiple embedded images (for ranges).
    """
    refs_found = (
        _IMAGE_REF.findall(text)
        + [m.group(0) for m in _IMAGE_REF_ALT.finditer(text)]
        + [m.group(0) for m in _IMAGE_REF_RANGE.finditer(text)]
    )
    paths = []
    for ref in refs_found:
        img_paths = ref_map.get(ref, [])
        if img_paths:
            for p in img_paths:
                paths.append((ref, p))
            text = text.replace(ref, f"[IMAGE:{ref}]")
    return text, paths


def _add_inline_content(paragraph, text: str, ref_paths: list, doc):
    """
    Add text content to a paragraph, inserting images as separate paragraphs
    appended directly to the document after the text paragraph.
    """
    from docx.shared import Inches, Pt

    # Add text with formatting (strip image placeholders for both ref formats)
    clean_text = re.sub(r'\[IMAGE:\[frame[^\]]+\]\]', '', text).strip()
    if clean_text:
        _add_formatted_text(paragraph, clean_text)

    # Append each image as its own paragraph in the document
    for ref_str, img_path in ref_paths:
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        try:
            run.add_picture(str(img_path), width=Inches(5.5))
            # Caption paragraph
            cap = doc.add_paragraph(ref_str.strip("[]"))
            for r in cap.runs:
                r.font.size = Pt(9)
                r.font.italic = True
        except Exception as exc:
            import logging
            logging.getLogger("exporter").warning(
                f"Could not embed image {ref_str}: {exc}"
            )
            for child in list(img_para._p):
                img_para._p.remove(child)
            img_para.add_run(f"[Image not embedded: {ref_str.strip('[]')}]")


def _add_formatted_text(paragraph, text: str):
    """Add text with basic bold/italic Markdown formatting."""
    # Handle **bold** and *italic*
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part:
            paragraph.add_run(part)


def _render_table(doc, table_lines: list):
    """Render a Markdown table as a DOCX table."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # Parse rows, skip separator line (---+---)
    rows = []
    for line in table_lines:
        if re.match(r"^\s*\|[-:\s|]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    # Normalize row lengths
    rows = [r + [""] * (max_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    from docx.shared import Inches
    col_width = Inches(6.5 / max_cols)

    for row_idx, row_data in enumerate(rows):
        tr = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = tr.cells[col_idx]
            cell.width = col_width
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if row_idx == 0:
                run.bold = True
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(10)

    doc.add_paragraph()


# ── Post-export image verification ───────────────────────────────────────────

def _count_embedded_images(docx_path: Path) -> int:
    """Return the number of images actually embedded in the DOCX."""
    from docx import Document
    doc = Document(str(docx_path))
    return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)


# ── Export recommendation logic ───────────────────────────────────────────────

def _should_recommend_docx(report_text: str) -> tuple[bool, str]:
    """Return (recommend_docx, reason)."""
    refs = _IMAGE_REF.findall(report_text)
    if len(refs) > 3:
        return True, f"Report contains {len(refs)} image references"
    return False, ""


# ── Main export flow ──────────────────────────────────────────────────────────

def run_export(
    meeting_folder: Path,
    output_format: str | None,
    report_path: Path | None = None,
):
    print("\n" + "═" * 56)
    print("  MeetingTool v2.0 — Export Report")
    print("═" * 56)

    if not meeting_folder.exists():
        _err(f"Meeting folder not found: {meeting_folder}")
        raise RuntimeError(f"Meeting folder not found: {meeting_folder}")

    # Use provided report path or auto-discover the most recent one
    if report_path is None:
        report_path = _find_report_md(meeting_folder)
    if not report_path or not report_path.exists():
        _err("No report_*.md found in this folder.")
        print("  Generate a report first by running 'mip run' and completing the LLM analysis.")
        raise RuntimeError("No report_*.md found in this folder")

    _ok(f"Report found: {report_path.name}")
    report_text = report_path.read_text(encoding="utf-8")

    # Locate frames directory
    frames_dir = meeting_folder / "imagenes_reunion"
    if not frames_dir.exists():
        _warn("imagenes_reunion\\ folder not found. Image embedding will be skipped.")
        frames_dir = None

    # Resolve image refs
    resolved = []
    missing  = []
    if frames_dir:
        resolved, missing = _resolve_image_refs(report_text, frames_dir)

    # Fail fast on missing refs
    if missing:
        _err(f"Cannot export: {len(missing)} image reference(s) not resolved.")
        for m in missing:
            print(f"    Missing: {m}")
            print(f"    Expected at: {frames_dir / m}")
        print()
        print("  Options:")
        print("    1. Run 'mip run' again to regenerate frames")
        print("    2. Remove the missing refs from report.md")
        raise RuntimeError(f"Cannot export: {len(missing)} image reference(s) not found in imagenes_reunion\\")

    # Recommend format if not specified
    if output_format is None:
        recommend_docx, reason = _should_recommend_docx(report_text)
        if recommend_docx:
            print(f"\n  {reason}. DOCX recommended for client delivery.")
        print(f"\n  Export format:")
        print(f"    [1] DOCX  — with embedded images (recommended for client delivery)")
        print(f"    [2] MD    — keep as Markdown only")
        print(f"    [3] Both")
        default = "1" if recommend_docx else "2"
        choice  = input(f"  Choice [{default}]: ").strip() or default
        format_map = {"1": "docx", "2": "md", "3": "both"}
        output_format = format_map.get(choice, "docx")

    # DOCX name matches the source .md (same stem, different extension)
    docx_path = meeting_folder / f"{report_path.stem}.docx"

    # Generate DOCX
    if output_format in ("docx", "both"):
        print(f"\n  Generating DOCX: {docx_path.name}")
        print(f"  Embedding {len(resolved)} image(s)...")
        try:
            _md_to_docx(report_text, resolved, docx_path, meeting_folder)
        except Exception as e:
            _err(f"DOCX generation failed: {e}")
            raise RuntimeError(f"DOCX generation failed: {e}") from e

        # Verify images are actually in the file
        actual = _count_embedded_images(docx_path)
        if resolved and actual == 0:
            _err(f"DOCX generated but contains 0 images — {len(resolved)} were expected.")
            raise RuntimeError(
                f"Export produced an empty DOCX (no images embedded). "
                f"Check that imagenes_reunion\\ exists and contains the frames."
            )
        elif actual < len(resolved):
            _warn(f"DOCX contains {actual} of {len(resolved)} expected images.")
        else:
            _ok(f"DOCX saved: {docx_path.name}  ({actual} image(s) embedded)")

    if output_format in ("md", "both"):
        _ok(f"Markdown report: {report_path.name} (already exists)")

    print()
    print("  Export complete.")
    if output_format in ("docx", "both"):
        print(f"  Ready for client delivery: {docx_path}")
    print()
